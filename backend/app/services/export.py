"""Commentary export — render an approved brief to PDF, Word, or email (.eml).

Rendering libraries (reportlab, python-docx) are imported lazily so the rest of
the API stays importable even if they are not installed; the endpoints surface a
clear 501 in that case.
"""

from __future__ import annotations

import io
import logging
from email.message import EmailMessage

from app.models.commentary import CompliantCommentary

logger = logging.getLogger(__name__)


class ExportUnavailableError(RuntimeError):
    """Raised when an optional rendering dependency is not installed."""


def _title(commentary: CompliantCommentary) -> str:
    return f"Portfolio Commentary — {commentary.mandate_id} — {commentary.period}"


def render_plaintext(commentary: CompliantCommentary) -> str:
    """Human-readable plain-text rendering (also the email body)."""
    lines: list[str] = [_title(commentary), ""]
    lines.append(f"Audience: {commentary.audience.value}")
    lines.append(f"Compliance status: {commentary.compliance_status.value}")
    lines.append("")
    for section in commentary.sections:
        lines.append(section.heading.value.upper())
        for claim in section.claims:
            value = f" ({claim.value})" if claim.value else ""
            lines.append(f"  - {claim.text}{value}")
        lines.append("")
    if commentary.disclaimers:
        lines.append("DISCLAIMERS")
        for d in commentary.disclaimers:
            lines.append(f"  - {d}")
        lines.append("")
    if commentary.source_map:
        lines.append("SOURCES")
        for sid, label in commentary.source_map.items():
            lines.append(f"  [{sid}] {label}")
    return "\n".join(lines)


def to_pdf(commentary: CompliantCommentary) -> bytes:
    try:
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ExportUnavailableError("reportlab is not installed (pip install reportlab).") from exc

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        leftMargin=20 * mm, rightMargin=20 * mm, topMargin=20 * mm, bottomMargin=20 * mm,
        title=_title(commentary),
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("wg-h1", parent=styles["Heading1"], fontSize=16, spaceAfter=10)
    h2 = ParagraphStyle("wg-h2", parent=styles["Heading2"], fontSize=12, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("wg-body", parent=styles["BodyText"], alignment=TA_LEFT, fontSize=10, leading=14)
    small = ParagraphStyle("wg-small", parent=body, fontSize=8, textColor="#555555")

    story: list = [Paragraph(_title(commentary), h1)]
    story.append(Paragraph(
        f"Audience: {commentary.audience.value} &nbsp;|&nbsp; "
        f"Compliance: {commentary.compliance_status.value}", small,
    ))
    story.append(Spacer(1, 6))

    for section in commentary.sections:
        story.append(Paragraph(section.heading.value, h2))
        items = []
        for claim in section.claims:
            value = f" <b>({claim.value})</b>" if claim.value else ""
            items.append(ListItem(Paragraph(f"{claim.text}{value}", body)))
        if items:
            story.append(ListFlowable(items, bulletType="bullet"))

    if commentary.disclaimers:
        story.append(Paragraph("Disclaimers", h2))
        story.append(ListFlowable(
            [ListItem(Paragraph(d, small)) for d in commentary.disclaimers], bulletType="bullet",
        ))
    buffer.seek(0)
    doc.build(story)
    return buffer.getvalue()


def to_docx(commentary: CompliantCommentary) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ExportUnavailableError("python-docx is not installed (pip install python-docx).") from exc

    document = Document()
    document.add_heading(_title(commentary), level=0)
    meta = document.add_paragraph()
    meta.add_run(
        f"Audience: {commentary.audience.value}  |  Compliance: {commentary.compliance_status.value}"
    ).italic = True

    for section in commentary.sections:
        document.add_heading(section.heading.value, level=1)
        for claim in section.claims:
            para = document.add_paragraph(style="List Bullet")
            para.add_run(claim.text)
            if claim.value:
                run = para.add_run(f" ({claim.value})")
                run.bold = True

    if commentary.disclaimers:
        document.add_heading("Disclaimers", level=1)
        for d in commentary.disclaimers:
            document.add_paragraph(d, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def to_eml(commentary: CompliantCommentary, to_address: str, from_address: str = "advisor@wealthgen.example") -> bytes:
    """Build an RFC 822 email with the brief as the body and a PDF attachment."""
    msg = EmailMessage()
    msg["Subject"] = _title(commentary)
    msg["From"] = from_address
    msg["To"] = to_address
    msg.set_content(render_plaintext(commentary))
    try:
        msg.add_attachment(
            to_pdf(commentary),
            maintype="application",
            subtype="pdf",
            filename=f"{commentary.mandate_id}_{commentary.period}.pdf",
        )
    except ExportUnavailableError:
        logger.info("PDF attachment skipped (reportlab unavailable); sending text-only email.")
    return msg.as_bytes()
